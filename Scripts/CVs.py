import os.path

import docx.document
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Inches

RIGHT_TAB_INCHES = 7
BULLET_POINT_STYLE = "ListBullet2"


def set_style_to_defaults(document: docx.document.Document, style_name: str):
    # Set the default styling
    style = document.styles[style_name]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = 0  # single

    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)


class InformationEntry:
    def __init__(self, title: str, description: str, start_date, end_date):
        self.title: str = title
        self.description: str = description
        self.start_date = start_date
        self.end_date = end_date


class Experience(InformationEntry):
    def __init__(self, title: str, description: str, company_name: str, location: str, start_date, end_date):
        super().__init__(title, description, start_date, end_date)
        self.company_name: str = company_name
        self.location: str = location


class Education(InformationEntry):
    def __init__(self, institution: str, location: str, degree: str, extracurriculars: str,
                 description: str, start_date, end_date):
        super().__init__(institution, description, start_date, end_date)
        self.location = location
        self.degree: str = degree
        self.extracurriculars: list[str] = extracurriculars.split(',')


class Project(InformationEntry):
    def __init__(self, title: str, description: str, start_date, end_date):
        super().__init__(title, description, start_date, end_date)


class CurriculumVitae:
    def __init__(self, name: str, experiences: list[Experience] = None, educations: list[Education] = None,
                 projects: list[Project] = None):
        super().__init__()
        self.document = Document()

        # Set the margins
        section = self.document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.left_margin = Inches(0.75)

        self.__right_tab_position = section.page_width - section.right_margin - section.left_margin

        # Set the default styling
        set_style_to_defaults(self.document, "Normal")
        set_style_to_defaults(self.document, BULLET_POINT_STYLE)

        # Set variables
        self.name = name
        self.experiences = experiences if experiences is not None else []
        self.educations = educations if educations is not None else []
        self.projects = projects if projects is not None else []

        # Create document
        self.__create_header__()
        self.__add_space__()
        self.__create_education__()
        self.__add_space__()
        self.__create_projects__()
        self.__add_space__()
        self.__create_experiences__()

    def __add_space__(self):
        self.document.add_paragraph()

    def __add_section_heading__(self, text: str):
        heading_run = self.document.add_paragraph(text).runs[0]
        heading_run.bold = True
        heading_run.font.size = Pt(12)

    def __create_header__(self):
        section = self.document.sections[0]
        section.different_first_page_header_footer = True
        header = section.first_page_header
        # Clear the existing paragraphs
        for paragraph in header.paragraphs:
            element = paragraph._element
            element.getparent().remove(element)

        # Add the name large at the top
        name_header = header.add_paragraph(self.name)
        name_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_header.runs[0]
        name_run.bold = True
        name_run.font.size = Pt(28)

        # Add a placeholder for the contact information
        contact_header = header.add_paragraph("[EMAIL] | [LINKEDIN] | [PHONE #]")
        contact_header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def __create_experiences__(self):
        if len(self.experiences) == 0:
            return
        self.__add_section_heading__("WORK EXPERIENCE")

        for experience in self.experiences:
            # Header for the position
            header_paragraph = self.document.add_paragraph()
            title_run = header_paragraph.add_run(experience.title.strip())
            title_run.bold = True

            company_and_location_run = header_paragraph.add_run(
                f" - {experience.company_name.strip()}, {experience.location.strip()}")
            company_and_location_run.italic = True

            tab_stops = header_paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(self.__right_tab_position, WD_TAB_ALIGNMENT.RIGHT)

            date_string = f"{experience.start_date.strip()} – {'Present' if experience.end_date == '' else experience.end_date.strip()}"
            date_run = header_paragraph.add_run()
            date_run.add_tab()
            date_run.add_text(date_string)
            date_run.italic = True

            # Any description of the position
            # TODO: Make this format nicer
            last_para = None
            for bullet in experience.description.split('•'):
                bullet.strip()
                if bullet == '':
                    continue
                last_para = self.document.add_paragraph(bullet, style=BULLET_POINT_STYLE)
            if last_para is not None:
                last_para.paragraph_format.space_after = Pt(2)

    def __create_education__(self):
        if len(self.educations) == 0:
            return
        self.__add_section_heading__("EDUCATION")

        for education in self.educations:
            # Header for the position
            header_paragraph = self.document.add_paragraph()
            title_run = header_paragraph.add_run(education.title.strip())
            title_run.bold = True

            company_and_location_run = header_paragraph.add_run(f" - {education.location.strip()}")
            company_and_location_run.italic = True

            tab_stops = header_paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(self.__right_tab_position, WD_TAB_ALIGNMENT.RIGHT)

            # TODO: Make this format to anticipated if it's not done
            date_string = f"{education.start_date.strip()} – {education.end_date.strip()}"
            date_run = header_paragraph.add_run()
            date_run.add_tab()
            date_run.add_text(date_string)
            date_run.italic = True

            # Add the degree / field of study
            self.document.add_paragraph(f"{education.degree.strip()}")

            # TODO: Minors

            # Extracurriculars
            if len(education.extracurriculars) > 0:
                ec_header = self.document.add_paragraph("Extracurriculars")
                ec_header.runs[0].italic = True
                for extracurricular in education.extracurriculars:
                    self.document.add_paragraph(extracurricular.strip(), style=BULLET_POINT_STYLE)

    def __create_projects__(self):
        if len(self.projects) == 0:
            return
        self.__add_section_heading__("PROJECTS")

        for project in self.projects:
            # Header for the position
            header_paragraph = self.document.add_paragraph()
            title_run = header_paragraph.add_run(project.title.strip())
            title_run.bold = True

            tab_stops = header_paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(self.__right_tab_position, WD_TAB_ALIGNMENT.RIGHT)

            date_string = f"{project.start_date.strip()} – {'Present' if project.end_date == '' else project.end_date.strip()}"
            date_run = header_paragraph.add_run()
            date_run.add_tab()
            date_run.add_text(date_string)
            date_run.italic = True

            # Any description of the project
            # TODO: Make this format nicer
            last_para = None
            for bullet in project.description.split('•'):
                bullet.strip()
                if bullet == '':
                    continue
                last_para = self.document.add_paragraph(bullet, style=BULLET_POINT_STYLE)
            if last_para is not None:
                last_para.paragraph_format.space_after = Pt(2)

    def save(self, path: str):
        self.document.save(path)
